import os  
import pandas as pd  
import random  
from docx import Document  

# Function to get grade from marks  
def get_grade(mark):  
    if pd.isna(mark):  # Check for NaN  
        return "N/A"  # Handling missing marks  
    if 90 <= mark <= 100:  
        return "A*"  
    elif 80 <= mark < 90:  
        return "A"  
    elif 70 <= mark < 80:  
        return "B"  
    elif 60 <= mark < 70:  
        return "C"  
    elif 50 <= mark < 60:  
        return "D"  
    elif 40 <= mark < 50:  
        return "E"  
    elif 30 <= mark < 40:  
        return "F"  
    else:  
        return "U"  

# Dictionary for random teacher remarks based on grades  
remarks_dict = {  
    "A*": ["Excellent performance!", "Outstanding work!", "Keep up the great effort!"],  
    "A": ["Great job!", "Very impressive!", "You've done well!"],  
    "B": ["Good progress!", "Nice work, but there’s room for improvement.", "Keep it up!"],  
    "C": ["Fair performance, consider more effort.", "You can do better, stay focused.", "Needs improvement."],  
    "D": ["More effort required.", "Consider asking for help with the subject.", "Put in a bit more practice."],  
    "E": ["Struggling in this subject, needs additional attention.", "Consider attending extra tutoring sessions."],  
    "F": ["Critical improvement needed.", "Seek immediate help from the teacher."],  
    "U": ["Ungraded, please see the teacher for details."]  
}  

# Function to generate report card for a single student  
def create_report_card(student_data, template_path, output_path):  
    doc = Document(template_path)  

    # Replace placeholders on the first page  
    student_name = student_data['Student Name'].iloc[0]  
    student_form = str(student_data['Form'].iloc[0])  

    for paragraph in doc.paragraphs:  
        if '<NAME>' in paragraph.text:  
            paragraph.text = paragraph.text.replace('<NAME>', student_name)  
        if '<FORM>' in paragraph.text:  
            paragraph.text = paragraph.text.replace('<FORM>', student_form)  

    # Fill the table with subject data  
    table = doc.tables[0]  # Assuming the first table is the one to fill  
    for row in student_data.itertuples():  
        subject = row.Subject  
        term_mark = row.Term_Mark  
        grade = get_grade(term_mark)  
        remarks = random.choice(remarks_dict[grade])  # Select a random remark based on grade  

        # Add a new row to the table  
        cells = table.add_row().cells  # Add row and get cells  
        cells[0].text = subject  
        cells[1].text = str(term_mark)  
        cells[2].text = grade  
        cells[3].text = remarks  

    # Save the document for this student  
    doc.save(output_path)  

# Function to process all students  
def generate_all_reports(data, template_path, output_folder):  
    # Ensure output folder exists  
    os.makedirs(output_folder, exist_ok=True)  

    # Group by 'Student Name' and generate reports  
    for student_name, student_data in data.groupby('Student Name'):  
        # Process each student's data  
        output_path = os.path.join(output_folder, f"{student_name.replace(' ', '_')}_report.docx")  
        try:  
            create_report_card(student_data, template_path, output_path)  
            print(f"Report generated for {student_name}.")  
        except Exception as e:  
            print(f"Error generating report for {student_name}: {e}")  

# Example usage  
if __name__ == "__main__":  
    # Combine Excel files from teachers into one dataframe  
    files = ['acc.xlsx', 'comp.xlsx', 'Geo.xlsx']  # Replace with actual file paths  
    combined_data = pd.concat([pd.read_excel(file) for file in files])  

    # Ensure data is sorted and filtered correctly  
    combined_data = combined_data.sort_values(by=['Student Name', 'Subject']).reset_index(drop=True)  

    template_path = "REPORT.docx"  # Path to your Word template  
    output_folder = "Reports"  # Folder where reports will be saved  

    # Generate reports for all students  
    generate_all_reports(combined_data, template_path, output_folder)